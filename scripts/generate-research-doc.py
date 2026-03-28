#!/usr/bin/env python3
"""Generate Getplace Location Intelligence Research Document (.docx)"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(BASE, "Getplace_Location_Intelligence_Research.docx")

# Colors
NAVY = RGBColor(0x1A, 0x1A, 0x2E)
BLUE = RGBColor(0x22, 0x6B, 0xE8)
DARK_BLUE = RGBColor(0x16, 0x4E, 0xA8)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professional styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, "226BE8")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            run.font.name = "Calibri"
            run.font.color.rgb = BLACK
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0F4FF")

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    doc.add_paragraph("")
    return table


def add_screenshot(doc, filename, caption, width=6.0):
    """Add a screenshot with caption."""
    path = os.path.join(BASE, filename)
    if not os.path.exists(path):
        p = doc.add_paragraph(f"[Screenshot: {filename} -- file not found]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MEDIUM_GRAY
    run.font.name = "Calibri"


def add_heading(doc, text, level=1):
    """Add a styled heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.color.rgb = NAVY
        elif level == 2:
            run.font.color.rgb = DARK_BLUE
        elif level == 3:
            run.font.color.rgb = BLUE


def add_body(doc, text):
    """Add body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_quote(doc, text, source=""):
    """Add a styled quote block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"\u201c{text}\u201d")
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_BLUE
    if source:
        run2 = p.add_run(f"\n\u2014 {source}")
        run2.font.size = Pt(9)
        run2.font.color.rgb = MEDIUM_GRAY
        run2.font.name = "Calibri"


def add_bullet(doc, text, bold_prefix=""):
    """Add a bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"
    return p


def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ==========================================
    # TITLE PAGE
    # ==========================================
    for _ in range(4):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GETPLACE")
    run.font.size = Pt(36)
    run.font.color.rgb = BLUE
    run.bold = True
    run.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Location Intelligence\nдля экспансии QSR")
    run.font.size = Pt(24)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"

    doc.add_paragraph("")

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run(
        "Как данные о людях, трафике и территориях\n"
        "превращаются в решения об открытии ресторанов"
    )
    run.font.size = Pt(14)
    run.font.color.rgb = MEDIUM_GRAY
    run.font.name = "Calibri"
    run.italic = True

    for _ in range(6):
        doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Март 2026 | Исследование рынка и платформы\nКонфиденциально")
    run.font.size = Pt(11)
    run.font.color.rgb = MEDIUM_GRAY
    run.font.name = "Calibri"

    doc.add_page_break()

    # ==========================================
    # TABLE OF CONTENTS
    # ==========================================
    add_heading(doc, "Содержание", level=1)
    toc_items = [
        "1. Введение: проблема экспансии QSR",
        "2. Рынок Location Intelligence",
        "3. Платформа Getplace: архитектура решения",
        "4. Система скоринга: как работает оценка локаций",
        "5. Кейс: Subway UK \u2014 белые пятна в насыщенном рынке",
        "6. Кейс: Greggs \u2014 drive-thru экспансия через анализ трафика",
        "7. Кейс: Nando's \u2014 premium-позиционирование и демографический fit",
        "8. Конкурентный анализ",
        "9. Бизнес-модель и Unit Economics",
        "10. Заключение и дорожная карта",
        "Приложение A. Источники данных",
        "Приложение B. Глоссарий",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = "Calibri"
        run.font.color.rgb = NAVY if item[0].isdigit() else DARK_BLUE
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ==========================================
    # 1. INTRODUCTION
    # ==========================================
    add_heading(doc, "1. Введение: проблема экспансии QSR", level=1)

    add_body(doc,
        "Индустрия быстрого обслуживания (Quick Service Restaurant, QSR) в Великобритании \u2014 "
        "это рынок объёмом 23,6 млрд фунтов стерлингов с 49 000+ точек продаж. "
        "Крупнейшие сети \u2014 Greggs (2 739), McDonald's (~1 493), Subway (2 081), "
        "Domino's (1 372), KFC (~1 000) \u2014 ежегодно открывают сотни новых локаций. "
        "Только в ближайшие 2\u20133 года ведущие сети планируют открыть 500+ точек."
    )

    add_heading(doc, "1.1 Цена ошибки", level=2)
    add_body(doc,
        "Открытие одной точки QSR стоит от 200 000 до 1 850 000 фунтов в зависимости от бренда. "
        "60% независимых ресторанов закрываются в первый год работы (NRA). "
        "38% владельцев ресторанов называют неудачную локацию и ограниченную парковку "
        "главными причинами финансовых трудностей (National Restaurant Association, 2023). "
        "Рестораны в зонах с высокой проходимостью живут в среднем 114 месяцев (9,5 лет), "
        "тогда как в зонах с низким трафиком \u2014 всего 66 месяцев (5,5 лет)."
    )

    add_styled_table(doc,
        ["Бренд", "Инвестиции в открытие (GBP)", "Франшиза"],
        [
            ["McDonald's", "350K \u2013 1.85M", "25% собственных средств"],
            ["KFC", "700K \u2013 3.7M", "45 000 франшиза"],
            ["Burger King", "300K \u2013 1.5M", "35\u201350K франшиза"],
            ["Domino's", "200K \u2013 350K", "25 000 франшиза"],
            ["Subway", "100K \u2013 250K", "10 000 франшиза"],
        ]
    )

    add_heading(doc, "1.2 Текущие методы выбора локации", level=2)
    add_body(doc,
        "Большинство сетей до сих пор используют комбинацию Excel-таблиц, консалтинговых отчётов "
        "(50\u2013200K фунтов за исследование) и экспертной интуиции. Средний цикл принятия решения \u2014 "
        "от 3 до 6 месяцев. При этом исследование McKinsey показывает, что компании, использующие "
        "AI и геопространственные технологии, улучшают качество выбора локаций на 10\u201320%."
    )

    add_quote(doc,
        "It's an arms race for data-driven site selection. These platforms integrate sales "
        "projections with local median incomes, mobile phone location data, and competitor "
        "footprints to yield a 'probabilistic forecast' of store success.",
        "Craig Mercer, Senior Analyst, McKinsey"
    )

    add_heading(doc, "1.3 Почему Location Intelligence \u2014 следующий этап", level=2)
    add_body(doc,
        "95% топ-менеджеров считают геопространственные данные критически важными для бизнеса (BCG). "
        "82% руководителей ресторанных компаний планируют увеличить инвестиции в AI (Deloitte, 2024). "
        "Применение аналитики на основе данных увеличивает удовлетворённость клиентов на 15\u201330%. "
        "Рынок Location Intelligence оценивается в $25 млрд (2025) и вырастет до $50+ млрд к 2030 году."
    )

    add_screenshot(doc, "screenshots/01-insights-overview.png",
        "Рис. 1. Платформа Getplace: карта Великобритании с 517 станциями и AI Strategic Analysis", 6.0)

    doc.add_page_break()

    # ==========================================
    # 2. MARKET
    # ==========================================
    add_heading(doc, "2. Рынок Location Intelligence", level=1)

    add_heading(doc, "2.1 Глобальный рынок (TAM/SAM/SOM)", level=2)
    add_body(doc,
        "Глобальный рынок location intelligence (LI) достиг $25 млрд в 2025 году. "
        "Консенсус-прогноз ведущих аналитических агентств \u2014 рост до $47\u201354 млрд к 2030 году "
        "с CAGR 13\u201317%. Специализированный сегмент геопространственной аналитики для ресторанов "
        "оценивается в $1,37 млрд (2024) и вырастет до $4,04 млрд к 2033 (CAGR 12,8%)."
    )

    add_styled_table(doc,
        ["Источник", "Прогноз 2030", "CAGR", "Период"],
        [
            ["Grand View Research", "$53.6B", "16.8%", "2025\u20132030"],
            ["Mordor Intelligence", "$47.1B", "13.5%", "2025\u20132030"],
            ["IMARC Group", "$68.8B (2033)", "13.1%", "2025\u20132033"],
            ["Precedence Research", "$74.8B (2035)", "16.0%", "2025\u20132035"],
        ],
        col_widths=[5, 3.5, 2.5, 3]
    )

    add_heading(doc, "2.2 Рынок QSR Великобритании", level=2)
    add_body(doc,
        "Рынок QSR в Великобритании \u2014 $33,8 млрд (2025), прогноз \u2014 $48,6 млрд к 2031 (CAGR 5,2%). "
        "49 055 точек быстрого питания, 468 000+ сотрудников. Формат travel hubs \u2014 самый быстрорастущий "
        "(CAGR 7,25% до 2031). Доставка (delivery) растёт ещё быстрее \u2014 6,71% CAGR."
    )

    add_heading(doc, "2.3 Ключевые тренды 2024\u20132026", level=2)
    add_bullet(doc, " AI-агенты в 40% корпоративных приложений к концу 2026 (Gartner, прогноз)", "AI-агенты:")
    add_bullet(doc, " Рынок геопространственного AI: $60 млрд (2025) \u2192 $472 млрд (2034), CAGR 25,8%", "GeoAI:")
    add_bullet(doc, " MCP (Model Context Protocol) от Anthropic становится стандартом для интеграции LLM с геоданными. "
        "Mapbox, CARTO, Foursquare уже выпустили MCP-серверы", "MCP:")
    add_bullet(doc, " Исследователи Penn State создали GIS Copilot с 86% успешных решений многошаговых задач. "
        "Foursquare + AWS развернули геопространственных AI-агентов в production", "Автономные GIS-агенты:")

    add_quote(doc,
        "By 2026, web search and geospatial tools will be tightly integrated into the major LLMs "
        "and assistants. Location and mapping will become built-in capabilities rather than add-ons.",
        "Mapbox Blog, GeoAI in 2026"
    )

    doc.add_page_break()

    # ==========================================
    # 3. PLATFORM
    # ==========================================
    add_heading(doc, "3. Платформа Getplace: архитектура решения", level=1)

    add_body(doc,
        "Getplace \u2014 insights-first платформа location intelligence для экспансии ресторанных сетей. "
        "В отличие от конкурентов, которые продают данные или инструменты, Getplace продаёт готовые инсайты: "
        "где открывать, почему, и с какой уверенностью. Ключевая фраза позиционирования: "
        "\u00abFind your next 50 locations before competitors do\u00bb."
    )

    add_heading(doc, "3.1 Три якоря анализа (Multi-Anchor Scoring)", level=2)
    add_body(doc,
        "Getplace использует уникальный подход multi-anchor scoring \u2014 три типа точек привязки, "
        "каждый со своей моделью оценки, объединённых в единый ранжированный список возможностей."
    )

    add_styled_table(doc,
        ["Тип якоря", "Модель", "Покрытие", "Ключевые сигналы"],
        [
            ["Станции (Stations)", "7 сигналов", "517 станций", "Footfall, brand gap, demographic, density, pedestrian, road traffic, workforce"],
            ["Перекрёстки (Junctions)", "4 сигнала", "5 000+ сегментов", "Traffic volume, drive-thru gap, QSR presence, demographic fit"],
            ["Зоны MSOA (Zones)", "4 сигнала", "1 100+ зон", "Workforce density, QSR gap, demographic fit, footfall proximity"],
        ],
        col_widths=[4, 3, 3.5, 6]
    )

    add_screenshot(doc, "screenshots/02-smart-map.png",
        "Рис. 2. Smart Map: единый интерфейс с картой и Intelligence Panel (6 630 возможностей)", 6.0)

    add_heading(doc, "3.2 7 AI-агентов", level=2)
    add_body(doc,
        "Система из 7 AI-агентов автоматически генерирует приоритизированные инсайты "
        "по 17 типам \u2014 от обнаружения белых пятен до мониторинга конкурентов."
    )

    add_styled_table(doc,
        ["Агент", "Стратегический вопрос", "Типы инсайтов"],
        [
            ["Market Monitor", "Что изменилось?", "Быстрый рост, смена лидеров, стагнация"],
            ["Competitor Tracker", "Кто угрожает?", "Доминирование, конкурентный вход, фланговая угроза"],
            ["Expansion Scout", "Где расти?", "Hot-opportunity, tier-upgrade, насыщение"],
            ["Delivery Intel", "Как доставлять?", "Покрытие платформ, delivery desert, drive-thru"],
            ["Human Flow Analyst", "Где люди?", "High-traffic/low-QSR станции, недообслуженные коридоры"],
            ["Market Fit Analyst", "Кто эти люди?", "Affluence mismatch, demographic expansion signal"],
            ["Opportunity Engine", "Где открывать?", "Конвергентная возможность (3+ сигнала)"],
        ],
        col_widths=[4, 4, 6]
    )

    add_screenshot(doc, "screenshots/06-insights-agents.png",
        "Рис. 3. Insights-вид с AI Strategic Analysis: автоматический анализ 517 станций", 6.0)

    add_heading(doc, "3.3 Источники данных", level=2)
    add_body(doc,
        "Getplace уникально интегрирует 8 источников данных: 7 государственных открытых датасетов "
        "(ORR, NaPTAN, DfT, IMD/WIMD/SIMD/NIMDM, Census 2021, BRES, ASHE) + собственные данные "
        "о 6 820+ ресторанных локациях 6 брендов. Это позволяет создавать инсайты, которые "
        "невозможно получить из одного источника."
    )

    add_quote(doc,
        "Ресторанная сеть может посмотреть на карту своих локаций. Но она не может "
        "наложить данные конкурентов + пассажиропоток + автобусные остановки + демографию и увидеть: "
        "\u00abУ вас 0 локаций рядом с Liverpool Street (98M пассажиров/год, 112 автобусных остановок), "
        "а McDonald's и KFC уже там \u2014 спрос подтверждён, ваша аудитория подходит, Score: 95/100.\u00bb "
        "Этот инсайт невозможен из одного источника. Он создаётся на пересечении восьми.",
        "Getplace: Human Traffic Intelligence"
    )

    doc.add_page_break()

    # ==========================================
    # 4. SCORING SYSTEM
    # ==========================================
    add_heading(doc, "4. Система скоринга: как работает оценка локаций", level=1)

    add_body(doc,
        "Каждый opportunity получает композитный скор от 0 до 100, рассчитанный как "
        "взвешенная сумма сигналов, умноженная на множитель уверенности (confidence). "
        "Минимум 2 сигнала должны \u00abсработать\u00bb для включения локации в список возможностей."
    )

    add_heading(doc, "4.1 Семь сигналов станций", level=2)

    add_styled_table(doc,
        ["#", "Сигнал", "Вес", "Срабатывает когда", "Обоснование"],
        [
            ["1", "Footfall (пассажиропоток)", "25%", ">1M пассажиров/год", "Корреляция 0.93 с продажами QSR (Unacast)"],
            ["2", "Brand Gap (бренд-разрыв)", "25%", "Бренд отсутствует, конкуренты есть", "Whitespace analysis \u2014 отраслевой стандарт"],
            ["3", "Demo Fit (демография)", "15%", "Income decile соответствует бренду", "40% потребителей <$50K сократили QSR (Restroworks)"],
            ["4", "Low Density (низкая плотность)", "15%", "QSR < 75% от среднего", "60% ресторанов закрываются в 1-й год (NRA)"],
            ["5", "Pedestrian (пешеходы)", "8%", "30+ автобусных остановок в 800м", "Пешеходы тратят на 40% больше (UW)"],
            ["6", "Road Traffic (дорожный трафик)", "7%", "50K+ авто/день, <2 drive-thru", "Drive-thru = 50\u201370% выручки QSR (US)"],
            ["7", "Workforce (рабочая сила)", "5%", "10K+ работников в 1.5км", "Стандартная метрика коммерческой недвижимости"],
        ],
        col_widths=[0.8, 4, 1.5, 4.5, 5]
    )

    add_screenshot(doc, "screenshots/04-station-detail.png",
        "Рис. 4. Station Deep Dive: Custom House \u2014 75/100, 7 сигналов, карта с 800м радиусом", 6.0)

    add_screenshot(doc, "screenshots/05-opportunity-cards.png",
        "Рис. 5. Карточки возможностей: Custom House (75), Liverpool Street (74), Bond Street (73), Waterloo (68)", 6.0)

    doc.add_page_break()

    # ==========================================
    # 5. CASE STUDY: SUBWAY
    # ==========================================
    add_heading(doc, "5. Кейс: Subway UK \u2014 белые пятна в насыщенном рынке", level=1)

    add_body(doc,
        "Subway \u2014 самый масштабный пример того, что происходит при экспансии без location intelligence. "
        "Сеть, достигнув пика в ~2 500 точек в Великобритании, потеряла 400\u2013500 локаций "
        "за последние годы из-за каннибализации и неудачного выбора мест."
    )

    add_heading(doc, "5.1 Масштаб проблемы", level=2)
    add_bullet(doc, " 2 081 локация в Великобритании (январь 2026) \u2014 все 100% франшизные", "Текущий статус:")
    add_bullet(doc, " С пика ~2 500 сеть потеряла ~18\u201320% точек", "Сокращение:")
    add_bullet(doc, " Глобально с 2015 года закрыто 7 600+ ресторанов (\u221228%)", "Глобальный тренд:")
    add_bullet(doc, " Каннибализация (точки слишком близко друг к другу), падение high street трафика, "
        "слабая unit economics франшизы", "Причины:")

    add_heading(doc, "5.2 Стратегический разворот", level=2)
    add_body(doc,
        "В августе 2023 Subway был продан Roark Capital за $9,6 млрд. Новая стратегия \u2014 "
        "\u00absmart growth\u00bb: правильная локация, правильный формат, правильный франчайзи. "
        "Планируется 10 000 новых ресторанов глобально, из них 2 500 в нетрадиционных "
        "локациях \u2014 ж/д станции, аэропорты, больницы, университеты."
    )

    add_heading(doc, "5.3 Как помогает Getplace", level=2)
    add_body(doc,
        "7-сигнальная модель Getplace напрямую адресует проблемы Subway: "
        "сигнал Brand Gap выявляет станции, где Subway отсутствует при наличии конкурентов. "
        "Сигнал Low Density предотвращает каннибализацию. "
        "82% точек Subway \u2014 в Англии; Шотландия, Уэльс и Северная Ирландия "
        "остаются недопредставленными \u2014 это видно в скоринге по регионам."
    )

    add_screenshot(doc, "screenshots/08-intelligence-brief.png",
        "Рис. 6. Executive Brief: сводка по 517 станциям, 5 000 перекрёсткам и 1 113 зонам", 6.0)

    doc.add_page_break()

    # ==========================================
    # 6. CASE STUDY: GREGGS
    # ==========================================
    add_heading(doc, "6. Кейс: Greggs \u2014 drive-thru экспансия через анализ трафика", level=1)

    add_body(doc,
        "Greggs \u2014 крупнейшая пекарная сеть Великобритании и пример агрессивной, "
        "но data-informed экспансии. С оборотом 2,2 млрд фунтов и долгосрочной целью "
        "в 3 500+ точек, Greggs нуждается в масштабном анализе трафика для drive-thru."
    )

    add_heading(doc, "6.1 Цифры роста", level=2)
    add_styled_table(doc,
        ["Показатель", "Значение"],
        [
            ["Текущие точки", "2 739 (2 137 собственных + 602 франшизных)"],
            ["Новые точки 2025", "121 net new (207 открытий, 36 закрытий)"],
            ["Drive-thru открыто", "45 (с 2017 года)"],
            ["Цель долгосрочная", "3 500+ точек"],
            ["Оборот 2025", "GBP 2.2 млрд (+6.8% YoY)"],
            ["Первый drive-thru", "Июнь 2017, Irlam Gateway, Salford"],
            ["Первый eco drive-thru", "Июнь 2025, Winchester"],
        ],
        col_widths=[5, 9]
    )

    add_heading(doc, "6.2 Стратегия выбора drive-thru локаций", level=2)
    add_body(doc,
        "Greggs целенаправленно размещает drive-thru на магистралях с высоким автомобильным "
        "трафиком \u2014 A-roads, retail parks, моторвейные сервисы (Welcome Break, Moto). "
        "Это напрямую соответствует 4-сигнальной junction-модели Getplace."
    )

    add_heading(doc, "6.3 Как помогает Getplace", level=2)
    add_body(doc,
        "Junction scoring Getplace анализирует 5 000+ высокотрафиковых дорожных сегментов (50K+ авто/день) "
        "по данным DfT AADF. Четыре сигнала \u2014 traffic volume, drive-thru gap, QSR presence, "
        "demographic fit \u2014 определяют перекрёстки с максимальным потенциалом для drive-thru. "
        "Diamond-маркеры на карте показывают возможности прямо в интерфейсе."
    )

    add_screenshot(doc, "screenshots/07-traffic-layer.png",
        "Рис. 7. Слой Road Traffic Flow: тепловая карта AADF + diamond-маркеры drive-thru возможностей", 6.0)

    doc.add_page_break()

    # ==========================================
    # 7. CASE STUDY: NANDO'S
    # ==========================================
    add_heading(doc, "7. Кейс: Nando's \u2014 premium-позиционирование и демографический fit", level=1)

    add_body(doc,
        "Nando's занимает уникальную нишу \u00abaffordable premium\u00bb \u2014 дороже QSR, "
        "дешевле casual dining. С выручкой 3,1 млн фунтов на ресторан (значительно выше QSR-среднего) "
        "Nando's показывает, что тщательный выбор локации окупается."
    )

    add_heading(doc, "7.1 Профиль бренда", level=2)
    add_styled_table(doc,
        ["Показатель", "Значение"],
        [
            ["Точки в Великобритании", "473 (100% company-owned)"],
            ["Новые точки/год", "12\u201314"],
            ["Выручка FY2025", "GBP 1.48 млрд (+8% YoY)"],
            ["Выручка на ресторан", "~GBP 3.1 млн/год"],
            ["Средний чек", "GBP 12\u201316 на человека"],
            ["Целевая аудитория", "Millennials/Gen Z, C1 socioeconomic group"],
            ["Модель", "100% company-owned (полный контроль)"],
        ],
        col_widths=[5, 9]
    )

    add_heading(doc, "7.2 Методология выбора локации", level=2)
    add_body(doc,
        "Nando's оценивает три ключевых фактора: (1) численность и плотность населения, "
        "(2) видимость и проходящий трафик (пешеходный + автомобильный), "
        "(3) соответствие района бренду \u2014 молодой, городской, динамичный. "
        "Premium-бренд требует income decile >= 6."
    )

    add_heading(doc, "7.3 Как помогает Getplace", level=2)
    add_body(doc,
        "MSOA zone scoring Getplace определяет зоны с высокой плотностью рабочих мест "
        "(workforce density) и подходящим демографическим профилем. "
        "Слой Income Level показывает расчётную зарплату по MSOA (на основе BRES + ASHE). "
        "Например, localIncomeDecile различает Bond Street (decile 8) и Brixton (decile 3) \u2014 "
        "критическая разница для premium-бренда."
    )

    add_screenshot(doc, "screenshots/05-opportunity-cards.png",
        "Рис. 8. Intelligence Panel: фильтр по станциям с AI-рекомендациями для каждой возможности", 6.0)

    doc.add_page_break()

    # ==========================================
    # 8. COMPETITIVE ANALYSIS
    # ==========================================
    add_heading(doc, "8. Конкурентный анализ", level=1)

    add_body(doc,
        "Рынок location intelligence для ресторанов включает несколько категорий игроков: "
        "платформы foot traffic (Placer.ai), вертикальные решения для QSR (SiteZeus), "
        "консалтинговые data-провайдеры (CACI, Experian) и геопространственные платформы (CARTO). "
        "Ни один не занимает позицию Getplace."
    )

    add_heading(doc, "8.1 Ландшафт конкурентов", level=2)

    add_styled_table(doc,
        ["Компания", "Фокус", "Гео", "Цена", "Funding", "AI-скоринг"],
        [
            ["Placer.ai", "Foot traffic", "US", "$1K+/мес", "$278M", "Аналитика"],
            ["SiteZeus", "AI site selection", "US", "Enterprise", "$15M", "Revenue forecast"],
            ["CACI (Acorn)", "Геодемография", "UK", "GBP 15K+/год", "Private", "Нет"],
            ["Experian", "Consumer data", "UK+global", "GBP 28K+", "Public (LSE)", "Нет"],
            ["Geoblink", "LI for retail", "EU", "Enterprise", "~$8M", "Predictive"],
            ["Buxton", "Customer analytics", "US", "Enterprise", "PE (PSG)", "Site scoring"],
            ["CARTO", "Geospatial platform", "Global", "$199+/мес", "$97M", "Toolkit"],
            ["Getplace", "QSR insights", "UK", "TBD", "Pre-seed", "7-signal + agents"],
        ],
        col_widths=[3, 3, 1.5, 3, 2.5, 3]
    )

    add_heading(doc, "8.2 Конкурентные преимущества Getplace", level=2)

    add_bullet(doc,
        " Ни один конкурент не интегрирует ORR (ж/д пассажиропоток), DfT (автотрафик), "
        "Census (рабочее население), IMD (депривация) и данные о брендах в единую модель скоринга. "
        "Placer.ai и SiteZeus ориентированы на US. CACI и Experian продают сырые данные, а не инсайты.",
        "UK-first + open data: "
    )
    add_bullet(doc,
        " Placer.ai \u2014 только foot traffic. CACI \u2014 только демография. Experian \u2014 только сегментация. "
        "Getplace объединяет 7 сигналов в единый скор 0\u2013100 с объяснением каждого сигнала.",
        "Multi-signal composite scoring: "
    )
    add_bullet(doc,
        " Уникальная рамка: станции + перекрёстки + зоны MSOA как три типа якорей "
        "с собственными моделями, объединённые в один ранжированный список.",
        "Multi-anchor framework: "
    )
    add_bullet(doc,
        " CARTO и Experian продают инструменты. Getplace показывает \u00abвот ваши топ-возможности\u00bb "
        "и объясняет \u00abпочему\u00bb через разбивку сигналов и AI-агентный анализ.",
        "Insights-first, не tools-first: "
    )
    add_bullet(doc,
        " Ни один конкурент не использует систему AI-агентов для контекстных рекомендаций "
        "и нарративного анализа каждой возможности. SiteZeus использует ML для прогнозирования, "
        "но не генерирует объяснительный intelligence.",
        "AI-agent orchestration: "
    )

    add_screenshot(doc, "screenshots/09-table-view.png",
        "Рис. 9. Table view: 12 регионов UK с разбивкой по 6 брендам, населению и плотности", 6.0)

    doc.add_page_break()

    # ==========================================
    # 9. BUSINESS MODEL
    # ==========================================
    add_heading(doc, "9. Бизнес-модель и Unit Economics", level=1)

    add_heading(doc, "9.1 Ценность для клиента", level=2)
    add_body(doc,
        "Прямая ценность: экономия 5 дней/мес аналиста (2 000 GBP/мес), "
        "повышение качества решений \u2014 при стоимости ошибки 200\u2013500K GBP "
        "даже 10% снижение риска = 20\u201350K GBP экономии на каждое решение. "
        "Стратегическая ценность: скорость получения инсайтов (\u00abоткрыл и посмотрел\u00bb "
        "вместо \u00abзаказал исследование, получил через месяц\u00bb)."
    )

    add_quote(doc,
        "One wrong location costs the client $500K\u2013$2M. A $30\u201350K/year subscription "
        "is a rounding error in their expansion budget.",
        "Getplace pricing rationale"
    )

    add_heading(doc, "9.2 Ценообразование", level=2)
    add_body(doc,
        "Целевой диапазон: $30\u201350K/год за подписку. Оси ценообразования: "
        "количество стран, брендов, глубина данных, количество пользователей, "
        "частота обновления. PLG-воронка: бесплатный Explorer (\u00abчто происходит?\u00bb) "
        "\u2192 платная платформа (\u00abчто делать?\u00bb)."
    )

    add_styled_table(doc,
        ["Бесплатно (Explorer)", "Платно (Platform)"],
        [
            ["Карта, базовый drill-down", "Полный Radar с настройкой весов"],
            ["Ограниченный набор брендов", "Все бренды + кастомные группы"],
            ["Timeline без аналитики", "Timeline с разбивкой + агенты"],
            ["Базовые метрики", "AI-агенты + алерты (email/Slack)"],
            ["\u2014", "Экспорт, API"],
        ],
        col_widths=[7, 7]
    )

    add_heading(doc, "9.3 Финансовые показатели", level=2)
    add_bullet(doc, " ~$500K", "Выручка 2025:")
    add_bullet(doc, " $1.5M (3x рост)", "Цель 2026:")
    add_bullet(doc, " 30% топ-50 глобальных ресторанных компаний, $1M/мес", "Цель 3 года:")
    add_bullet(doc, " Burger King (UK + Germany)", "Текущий клиент:")
    add_bullet(doc, " <10 человек", "Команда:")
    add_bullet(doc, " Expand within clients (BK = 2 из 100+ стран), reference selling, Explorer как воронка", "Стратегия роста:")

    add_heading(doc, "9.4 TAM для UK QSR", level=2)
    add_body(doc,
        "В Великобритании ~50 крупных ресторанных сетей с потребностью в location intelligence. "
        "При цене $30\u201350K/год: SAM = $1.5\u20132.5M только на UK. "
        "Глобально топ-50 ресторанных компаний \u2014 TAM = $50\u2013100M+ "
        "(при $100\u2013200K/год на multi-country подписку)."
    )

    add_screenshot(doc, "screenshots/10-map-london.png",
        "Рис. 10. Map view: станции со слоем Station Analysis + Deprivation Index", 6.0)

    doc.add_page_break()

    # ==========================================
    # 10. CONCLUSION
    # ==========================================
    add_heading(doc, "10. Заключение и дорожная карта", level=1)

    add_heading(doc, "10.1 Реализовано", level=2)
    add_body(doc,
        "Getplace прошёл 5 фаз разработки: MVP \u2192 Depth \u2192 Dynamics \u2192 Scale \u2192 Intelligence. "
        "На текущий момент реализовано:"
    )
    add_bullet(doc, " 6 820 реальных ресторанных локаций 6 брендов", "")
    add_bullet(doc, " 2 361 станция с 7-сигнальным скорингом", "")
    add_bullet(doc, " 5 000 точек измерения трафика", "")
    add_bullet(doc, " 43 535 микро-областей с нормализованной демографией", "")
    add_bullet(doc, " 7 AI-агентов, генерирующих ~46 инсайтов", "")
    add_bullet(doc, " Smart Map с multi-anchor Intelligence Panel", "")
    add_bullet(doc, " 17,6 млн расчётов proximity (station-restaurant) через Python ETL", "")

    add_heading(doc, "10.2 Следующие шаги", level=2)
    add_bullet(doc, " Выход за пределы UK (Германия, Франция, Испания, Италия)", "Мультистрановость:")
    add_bullet(doc, " REST API для интеграции с BI-платформами (Tableau, Looker)", "API:")
    add_bullet(doc, " Переход от rule-based к LLM-powered агентам", "LLM-агенты:")
    add_bullet(doc, " MCP-сервер для интеграции с Claude, Cursor и другими AI-инструментами", "MCP-интеграция:")
    add_bullet(doc, " Кросс-страновое сравнение и белые пятна глобальных сетей", "Cross-country:")

    add_heading(doc, "10.3 Видение", level=2)
    add_body(doc,
        "Getplace стремится стать стандартом location intelligence для ресторанной индустрии: "
        "платформа, которую используют 30% топ-50 мировых ресторанных компаний для принятия "
        "решений об экспансии. Не инструмент для работы с данными, а источник готовых инсайтов \u2014 "
        "\u00abWe sell not data, but the interpretation of data.\u00bb"
    )

    add_screenshot(doc, "screenshots/11-radar.png",
        "Рис. 11. Expansion Radar: 12 регионов UK ранжированы по потенциалу экспансии для Subway", 5.5)

    doc.add_page_break()

    # ==========================================
    # APPENDIX A: DATA SOURCES
    # ==========================================
    add_heading(doc, "Приложение A. Источники данных платформы", level=1)

    add_styled_table(doc,
        ["Источник", "Издатель", "Что даёт", "Покрытие"],
        [
            ["ORR Station Usage", "Office of Rail and Road", "Ежегодные входы/выходы на каждой станции UK", "~2 600 станций"],
            ["NaPTAN", "DfT", "GPS-координаты всех остановок общественного транспорта", "~434K записей"],
            ["DfT Traffic Counts", "DfT", "Среднегодовой автомобильный поток (AADF)", "~46 000 точек"],
            ["IMD / WIMD / SIMD / NIMDM", "MHCLG / Welsh Gov / Scottish Gov / NISRA", "Индексы депривации", "~43 535 микро-областей"],
            ["Census 2021 (WP001)", "ONS", "Рабочее население по MSOA", "7 264 зоны"],
            ["BRES", "ONS", "Занятость по отраслям (SIC) и районам", "348 LA x 19 отраслей"],
            ["ASHE", "ONS", "Медианная зарплата по отраслям", "19 отраслей SIC"],
            ["Getplace (собственные)", "Getplace", "6 820+ ресторанных локаций 6 брендов", "UK-wide"],
        ],
        col_widths=[4, 3.5, 5, 3]
    )

    add_screenshot(doc, "screenshots/12-deprivation-layer.png",
        "Рис. 12. Слой Deprivation Index: MSOA-уровень, данные IMD 2025 + WIMD 2025", 5.5)

    doc.add_page_break()

    # ==========================================
    # APPENDIX B: GLOSSARY
    # ==========================================
    add_heading(doc, "Приложение B. Глоссарий", level=1)

    add_styled_table(doc,
        ["Термин", "Определение"],
        [
            ["QSR (Quick Service Restaurant)", "Формат быстрого обслуживания: McDonald's, KFC, Subway и т.д."],
            ["AADF (Annual Average Daily Flow)", "Среднегодовой суточный поток автотранспорта на участке дороги"],
            ["MSOA (Middle Super Output Area)", "Статистическая зона среднего уровня (~7 800 человек) в England & Wales"],
            ["IMD (Index of Multiple Deprivation)", "Индекс множественной депривации \u2014 комплексная оценка бедности района"],
            ["ORR (Office of Rail and Road)", "Регулятор железнодорожного транспорта Великобритании"],
            ["NaPTAN", "National Public Transport Access Nodes \u2014 все остановки общ. транспорта UK"],
            ["BRES", "Business Register and Employment Survey \u2014 занятость по отраслям и районам"],
            ["ASHE", "Annual Survey of Hours and Earnings \u2014 обзор зарплат по отраслям"],
            ["TAM / SAM / SOM", "Total/Serviceable/Obtainable Addressable Market \u2014 уровни оценки рынка"],
            ["PLG (Product-Led Growth)", "Стратегия роста через бесплатный продукт как воронку"],
            ["Multi-Anchor Scoring", "Подход Getplace: оценка через 3 типа якорей (станции, перекрёстки, зоны)"],
            ["Footfall", "Пешеходный трафик / пассажиропоток"],
            ["Brand Gap", "Отсутствие бренда в локации при наличии конкурентов"],
            ["Drive-thru", "Формат обслуживания из автомобиля"],
            ["Income Decile", "Дециль дохода: 1 = беднейшие 10%, 10 = богатейшие 10%"],
        ],
        col_widths=[5, 11]
    )

    # ==========================================
    # FOOTER
    # ==========================================
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("\u2014\u2014\u2014\n\nGetplace | Location Intelligence для QSR\nМарт 2026 | Конфиденциально")
    run.font.size = Pt(9)
    run.font.color.rgb = MEDIUM_GRAY
    run.font.name = "Calibri"

    # Save
    doc.save(OUT)
    print(f"Document saved: {OUT}")
    print(f"Size: {os.path.getsize(OUT) / 1024 / 1024:.1f} MB")


if __name__ == "__main__":
    build_document()
